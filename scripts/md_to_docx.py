"""
Utility to convert a Markdown document into DOCX.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document


def process_line(doc: Document, line: str) -> None:
    stripped = line.rstrip()
    if stripped == "":
        doc.add_paragraph("")
        return

    heading = re.match(r"^(#+)\s+(.*)$", stripped)
    if heading:
        level = min(len(heading.group(1)), 4)
        doc.add_heading(heading.group(2).strip(), level=level)
        return

    if stripped.startswith(("- ", "* ")):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return

    if re.match(r"^\d+\.\s+", stripped):
        doc.add_paragraph(stripped, style="List Number")
        return

    doc.add_paragraph(stripped)


def convert(markdown_path: Path, output_path: Path) -> None:
    if not markdown_path.exists():
        raise FileNotFoundError(markdown_path)

    document = Document()
    content = markdown_path.read_text(encoding="utf-8").splitlines()
    for line in content:
        process_line(document, line)
    document.save(output_path)


if __name__ == "__main__":
    project_root = Path(__file__).resolve().parents[1]
    md = project_root / "docs" / "PROJECT_DOCUMENTATION.md"
    docx = project_root / "docs" / "PROJECT_DOCUMENTATION.docx"
    convert(md, docx)
    print(f"Generated DOCX at {docx}")



